# -*- coding: utf-8 -*-
"""Generate editable and fixed-layout reports from the native report view."""

from __future__ import annotations

import html
import os
import re
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    ListFlowable,
    ListItem,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)


DOCX_BODY_FONT = "Arial Unicode MS"
PDF_CJK_FONT = "AdarianCJK"
INK = "17324D"
HEADING = "2E74B5"
HEADING_DARK = "1F4D78"
MUTED = "607586"
LIGHT_FILL = "F2F4F7"


def write_report_docx(view: dict[str, Any], path: Path) -> Path:
    """Write an editable Word report using the standard business brief preset."""

    path.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    _configure_docx(document)
    document.core_properties.title = str(view.get("title") or "舆情风险研判报告")

    title = document.add_paragraph()
    title.paragraph_format.space_after = Pt(6)
    title_run = title.add_run(_display_text(view.get("title") or "舆情风险研判报告"))
    _format_docx_run(title_run, size=22, bold=True, color=INK)

    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(6)
    subtitle_run = subtitle.add_run(_display_text(view.get("subtitle") or ""))
    _format_docx_run(subtitle_run, size=11, color=MUTED)

    metadata = document.add_paragraph()
    metadata.paragraph_format.space_after = Pt(14)
    source = view.get("source") or {}
    metadata_text = "生成时间：{}    Batch：{}    版本：{} 版".format(
        view.get("generated_at") or "",
        source.get("batch_id") or view.get("batch_id") or "",
        view.get("version") or "",
    )
    metadata_run = metadata.add_run(metadata_text)
    _format_docx_run(metadata_run, size=9, color=MUTED)

    _add_docx_kpis(document, view.get("kpis") or [])
    for section in view.get("sections") or []:
        _add_docx_section(document, section)
    _add_docx_appendix(document, view.get("appendix") or {})

    document.save(path)
    return path


def write_report_pdf(view: dict[str, Any], path: Path) -> Path:
    """Write a fixed-layout PDF report directly from the native report view."""

    path.parent.mkdir(parents=True, exist_ok=True)
    _register_pdf_font()

    styles = _pdf_styles()
    story: list[Any] = [
        Paragraph(_pdf_text(view.get("title") or "舆情风险研判报告"), styles["Title"]),
        Paragraph(_pdf_text(view.get("subtitle") or ""), styles["Subtitle"]),
    ]
    source = view.get("source") or {}
    metadata = "生成时间：{}　Batch：{}　版本：{} 版".format(
        view.get("generated_at") or "",
        source.get("batch_id") or view.get("batch_id") or "",
        view.get("version") or "",
    )
    story.extend([Paragraph(_pdf_text(metadata), styles["Meta"]), Spacer(1, 10)])
    _add_pdf_kpis(story, view.get("kpis") or [], styles)
    for section in view.get("sections") or []:
        _add_pdf_section(story, section, styles)
    _add_pdf_appendix(story, view.get("appendix") or {}, styles)

    report = SimpleDocTemplate(
        str(path),
        pagesize=letter,
        leftMargin=inch,
        rightMargin=inch,
        topMargin=0.78 * inch,
        bottomMargin=0.72 * inch,
        title=str(view.get("title") or "舆情风险研判报告"),
        author="Adarian",
    )
    report.build(story, onFirstPage=_draw_pdf_footer, onLaterPages=_draw_pdf_footer)
    return path


def _configure_docx(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.start_type = WD_SECTION.NEW_PAGE

    normal = document.styles["Normal"]
    _configure_docx_style(normal, 11, INK, 0, 6, 1.1)
    _configure_docx_style(document.styles["Heading 1"], 16, HEADING, 16, 8, 1.0, bold=True)
    _configure_docx_style(document.styles["Heading 2"], 13, HEADING, 12, 6, 1.0, bold=True)
    _configure_docx_style(document.styles["Heading 3"], 12, HEADING_DARK, 8, 4, 1.0, bold=True)

    list_style = document.styles["List Bullet"]
    _configure_docx_style(list_style, 11, INK, 0, 8, 1.167)
    list_style.paragraph_format.left_indent = Inches(0.5)
    list_style.paragraph_format.first_line_indent = Inches(-0.25)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_run = footer.add_run("Adarian  |  ")
    _format_docx_run(footer_run, size=9, color=MUTED)
    _add_page_field(footer)


def _configure_docx_style(style: Any, size: float, color: str, before: float, after: float, line_spacing: float, *, bold: bool = False) -> None:
    style.font.name = DOCX_BODY_FONT
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = RGBColor.from_string(color)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), DOCX_BODY_FONT)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.line_spacing = line_spacing


def _format_docx_run(run: Any, *, size: float = 11, bold: bool = False, color: str = INK) -> None:
    run.font.name = DOCX_BODY_FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), DOCX_BODY_FONT)


def _add_docx_kpis(document: Document, kpis: list[dict[str, Any]]) -> None:
    if not kpis:
        return
    table = document.add_table(rows=1, cols=len(kpis))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    width_dxa = 9360 // len(kpis)
    _set_table_geometry(table, [width_dxa] * len(kpis))
    for cell, kpi in zip(table.rows[0].cells, kpis):
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        _set_cell_shading(cell, LIGHT_FILL)
        _set_cell_margins(cell, top=120, bottom=120, start=120, end=120)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        label = paragraph.add_run(_display_text(kpi.get("label") or ""))
        _format_docx_run(label, size=8.5, color=MUTED)
        value = paragraph.add_run(f"\n{_display_text(kpi.get('value') or '')}\n")
        _format_docx_run(value, size=15, bold=True, color=INK)
        note = paragraph.add_run(_display_text(kpi.get("note") or ""))
        _format_docx_run(note, size=8, color=MUTED)
    document.add_paragraph().paragraph_format.space_after = Pt(2)


def _add_docx_section(document: Document, section: dict[str, Any]) -> None:
    document.add_heading(_display_text(section.get("heading") or ""), level=1)
    for block in section.get("blocks") or []:
        kind = block.get("type")
        if kind == "subheading":
            document.add_heading(_display_text(block.get("text") or ""), level=2)
        elif kind == "list":
            for item in block.get("items") or []:
                paragraph = document.add_paragraph(style="List Bullet")
                paragraph.add_run(_display_text(item))
        elif kind == "callout":
            paragraph = document.add_paragraph()
            _set_paragraph_shading(paragraph, "EEF8FC")
            title = paragraph.add_run(f"{_display_text(block.get('title') or '')}\n")
            _format_docx_run(title, bold=True, color=HEADING_DARK)
            paragraph.add_run(_display_text(block.get("text") or ""))
        else:
            text = block.get("text") or ""
            if _is_markdown_subheading(text):
                document.add_heading(_display_text(text), level=2)
            else:
                document.add_paragraph(_display_text(text))


def _add_docx_appendix(document: Document, appendix: dict[str, Any]) -> None:
    if appendix.get("mode") == "hidden":
        return
    document.add_heading("五、附录引用", level=1)
    document.add_paragraph(
        "事件：{}；completed worlds：{}；确认风险：{}；等级分布：{}".format(
            appendix.get("event_name") or "",
            appendix.get("worlds_count") or 0,
            appendix.get("confirmed_risks") or 0,
            appendix.get("risk_distribution") or "暂无",
        )
    )
    if appendix.get("mode") == "references":
        for reference in appendix.get("references") or []:
            paragraph = document.add_paragraph(style="List Bullet")
            paragraph.add_run(_display_text(reference))


def _set_table_geometry(table: Any, widths: list[int]) -> None:
    table_xml = table._tbl
    properties = table_xml.tblPr
    width = properties.first_child_found_in("w:tblW")
    if width is None:
        width = OxmlElement("w:tblW")
        properties.append(width)
    width.set(qn("w:type"), "dxa")
    width.set(qn("w:w"), str(sum(widths)))
    indent = OxmlElement("w:tblInd")
    indent.set(qn("w:type"), "dxa")
    indent.set(qn("w:w"), "120")
    properties.append(indent)
    for grid_col, value in zip(table_xml.tblGrid.gridCol_lst, widths):
        grid_col.set(qn("w:w"), str(value))
    for row in table.rows:
        for cell, value in zip(row.cells, widths):
            cell.width = Inches(value / 1440)
            cell._tc.tcPr.tcW.set(qn("w:type"), "dxa")
            cell._tc.tcPr.tcW.set(qn("w:w"), str(value))


def _set_cell_shading(cell: Any, fill: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shading)


def _set_cell_margins(cell: Any, **margins: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in margins.items():
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_paragraph_shading(paragraph: Any, fill: str) -> None:
    properties = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    properties.append(shading)


def _add_page_field(paragraph: Any) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, end])
    _format_docx_run(run, size=9, color=MUTED)


def _pdf_styles() -> dict[str, ParagraphStyle]:
    base = getSampleStyleSheet()
    return {
        "Title": ParagraphStyle(
            "AdarianTitle",
            parent=base["Title"],
            fontName=PDF_CJK_FONT,
            fontSize=22,
            leading=28,
            textColor=colors.HexColor(f"#{INK}"),
            alignment=TA_LEFT,
            spaceAfter=8,
            wordWrap="CJK",
        ),
        "Subtitle": ParagraphStyle(
            "AdarianSubtitle",
            parent=base["Normal"],
            fontName=PDF_CJK_FONT,
            fontSize=10.5,
            leading=16,
            textColor=colors.HexColor(f"#{MUTED}"),
            spaceAfter=6,
            wordWrap="CJK",
        ),
        "Meta": ParagraphStyle(
            "AdarianMeta",
            parent=base["Normal"],
            fontName=PDF_CJK_FONT,
            fontSize=8.5,
            leading=12,
            textColor=colors.HexColor(f"#{MUTED}"),
            wordWrap="CJK",
        ),
        "H1": ParagraphStyle(
            "AdarianH1",
            parent=base["Heading1"],
            fontName=PDF_CJK_FONT,
            fontSize=16,
            leading=22,
            textColor=colors.HexColor(f"#{HEADING}"),
            spaceBefore=16,
            spaceAfter=8,
            wordWrap="CJK",
        ),
        "H2": ParagraphStyle(
            "AdarianH2",
            parent=base["Heading2"],
            fontName=PDF_CJK_FONT,
            fontSize=12,
            leading=18,
            textColor=colors.HexColor(f"#{HEADING_DARK}"),
            spaceBefore=10,
            spaceAfter=5,
            wordWrap="CJK",
        ),
        "Body": ParagraphStyle(
            "AdarianBody",
            parent=base["BodyText"],
            fontName=PDF_CJK_FONT,
            fontSize=10.5,
            leading=17,
            textColor=colors.HexColor(f"#{INK}"),
            spaceAfter=7,
            wordWrap="CJK",
        ),
        "List": ParagraphStyle(
            "AdarianList",
            parent=base["BodyText"],
            fontName=PDF_CJK_FONT,
            fontSize=10.5,
            leading=17,
            textColor=colors.HexColor(f"#{INK}"),
            leftIndent=4,
            spaceAfter=4,
            wordWrap="CJK",
        ),
        "Kpi": ParagraphStyle(
            "AdarianKpi",
            parent=base["Normal"],
            fontName=PDF_CJK_FONT,
            fontSize=8,
            leading=12,
            alignment=TA_CENTER,
            textColor=colors.HexColor(f"#{MUTED}"),
            wordWrap="CJK",
        ),
    }


def _add_pdf_kpis(story: list[Any], kpis: list[dict[str, Any]], styles: dict[str, ParagraphStyle]) -> None:
    if not kpis:
        return
    cells = []
    for kpi in kpis:
        content = "{}<br/><font size=\"15\" color=\"#{}\">{}</font><br/><font size=\"7\">{}</font>".format(
            _pdf_text(kpi.get("label") or ""),
            INK,
            _pdf_text(kpi.get("value") or ""),
            _pdf_text(kpi.get("note") or ""),
        )
        cells.append(Paragraph(content, styles["Kpi"]))
    width = 6.5 * inch / len(cells)
    table = Table([cells], colWidths=[width] * len(cells), hAlign="CENTER")
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor(f"#{LIGHT_FILL}")),
        ("BOX", (0, 0), (-1, -1), 0.5, colors.HexColor("#DCE8EF")),
        ("INNERGRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#DCE8EF")),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("LEFTPADDING", (0, 0), (-1, -1), 8),
        ("RIGHTPADDING", (0, 0), (-1, -1), 8),
        ("TOPPADDING", (0, 0), (-1, -1), 9),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 9),
    ]))
    story.extend([table, Spacer(1, 8)])


def _add_pdf_section(story: list[Any], section: dict[str, Any], styles: dict[str, ParagraphStyle]) -> None:
    story.append(Paragraph(_pdf_text(section.get("heading") or ""), styles["H1"]))
    for block in section.get("blocks") or []:
        kind = block.get("type")
        if kind == "subheading":
            story.append(Paragraph(_pdf_text(block.get("text") or ""), styles["H2"]))
        elif kind == "list":
            items = [
                ListItem(Paragraph(_pdf_text(item), styles["List"]), leftIndent=12)
                for item in block.get("items") or []
            ]
            story.append(ListFlowable(items, bulletType="bullet", bulletFontName=PDF_CJK_FONT, leftIndent=20, bulletOffsetY=1))
            story.append(Spacer(1, 4))
        elif kind == "callout":
            title = _pdf_text(block.get("title") or "")
            text = _pdf_text(block.get("text") or "")
            callout = Table(
                [[Paragraph(f"<b>{title}</b><br/>{text}", styles["Body"])]],
                colWidths=[6.35 * inch],
            )
            callout.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#EEF8FC")),
                ("BOX", (0, 0), (-1, -1), 0.8, colors.HexColor("#A8D5EA")),
                ("LEFTPADDING", (0, 0), (-1, -1), 10),
                ("RIGHTPADDING", (0, 0), (-1, -1), 10),
                ("TOPPADDING", (0, 0), (-1, -1), 8),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
            ]))
            story.extend([callout, Spacer(1, 7)])
        else:
            text = block.get("text") or ""
            style = styles["H2"] if _is_markdown_subheading(text) else styles["Body"]
            story.append(Paragraph(_pdf_text(text), style))


def _add_pdf_appendix(story: list[Any], appendix: dict[str, Any], styles: dict[str, ParagraphStyle]) -> None:
    if appendix.get("mode") == "hidden":
        return
    story.append(Paragraph("五、附录引用", styles["H1"]))
    summary = "事件：{}；completed worlds：{}；确认风险：{}；等级分布：{}".format(
        appendix.get("event_name") or "",
        appendix.get("worlds_count") or 0,
        appendix.get("confirmed_risks") or 0,
        appendix.get("risk_distribution") or "暂无",
    )
    story.append(Paragraph(_pdf_text(summary), styles["Body"]))
    if appendix.get("mode") == "references":
        items = [
            ListItem(Paragraph(_pdf_text(item), styles["List"]), leftIndent=12)
            for item in appendix.get("references") or []
        ]
        story.append(ListFlowable(items, bulletType="bullet", bulletFontName=PDF_CJK_FONT, leftIndent=20, bulletOffsetY=1))


def _draw_pdf_footer(canvas: Any, document: Any) -> None:
    canvas.saveState()
    canvas.setFont(PDF_CJK_FONT, 8)
    canvas.setFillColor(colors.HexColor(f"#{MUTED}"))
    canvas.drawRightString(letter[0] - inch, 0.42 * inch, f"Adarian  |  {document.page}")
    canvas.restoreState()


def _pdf_text(value: Any) -> str:
    return html.escape(_display_text(value)).replace("\n", "<br/>")


def _display_text(value: Any) -> str:
    text = str(value or "").strip()
    text = re.sub(r"^#{1,6}\s+", "", text)
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"__(.+?)__", r"\1", text)
    text = re.sub(r"`([^`]+)`", r"\1", text)
    return text


def _is_markdown_subheading(value: Any) -> bool:
    return bool(re.match(r"^\s*#{2,6}\s+", str(value or "")))


def _register_pdf_font() -> None:
    if PDF_CJK_FONT in pdfmetrics.getRegisteredFontNames():
        return
    candidates = [
        os.getenv("ADARIAN_REPORT_FONT_PATH", ""),
        "/System/Library/Fonts/Supplemental/Arial Unicode.ttf",
        "/Library/Fonts/Arial Unicode.ttf",
        "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
        "/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc",
        "/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc",
    ]
    font_path = next((Path(item) for item in candidates if item and Path(item).is_file()), None)
    if font_path is None:
        raise RuntimeError(
            "REPORT_PDF_FONT_NOT_FOUND: configure ADARIAN_REPORT_FONT_PATH with a Chinese TrueType font"
        )
    pdfmetrics.registerFont(TTFont(PDF_CJK_FONT, str(font_path), subfontIndex=0))
